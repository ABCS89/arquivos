import pandas as pd
from docx import Document
from datetime import datetime
from num2words import num2words
import os

# Configurações
planilha = "teste.ods"
modelo = "modelo.docx"
saida = "cartas_geradas"

os.makedirs(saida, exist_ok=True)

# Lê a planilha
df = pd.read_excel(planilha, engine="odf")

# Campos adicionais (que podem vir do PDF ou ser fixos)
data_hoje = datetime.now().strftime("%d de %B de %Y")

# Função auxiliar para converter valor em extenso
def valor_por_extenso(valor):
    try:
        valor_float = float(str(valor).replace(",", "."))
        return num2words(valor_float, lang="pt_BR").replace(" e zero centavos", "")
    except:
        return str(valor)

# Gera cada documento
for _, row in df.iterrows():
    doc = Document(modelo)
    substituicoes = {
        "[nome]": str(row["Funcionário"]).title(),
        "[numero_funcional]": str(row["Nº Funcional"]),
        "[mensalidade]": str(row["Mensalidade"]),
        "[coparticipacao]": str(row["Coparticipação"]),
        "[valor_total]": str(row["Total"]),
        "[valor_extenso]": valor_por_extenso(row["Total"]),
        "[data_hoje]": data_hoje,
        # Se tiver PDF de e-mail com data/envio:
        "[data_envio_email]": "10 de outubro de 2025",  # exemplo
        "[email]": "email@exemplo.com",                 # exemplo
    }

    # Substitui nos parágrafos
    for paragrafo in doc.paragraphs:
        for chave, valor in substituicoes.items():
            if chave in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(chave, valor)

    # Substitui também nas tabelas (caso tenha)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for chave, valor in substituicoes.items():
                    celula.text = celula.text.replace(chave, valor)

    nome_arquivo = f"{row['Nº Funcional']}_{row['Funcionário'].replace(' ', '_')}.docx"
    doc.save(os.path.join(saida, nome_arquivo))

print("✅ Cartas geradas com sucesso!")
