import pandas as pd
from docx import Document
from docx.shared import Pt
import re
import os
import PyPDF2
from num2words import num2words
from datetime import datetime
import calendar

# ===============================
# CONFIGURAÇÕES
# ===============================
ods_path = 'template/teste.ods'
pdf_directory = 'template'
output_dir = 'output'

os.makedirs(output_dir, exist_ok=True)

# ===============================
# MESES
# ===============================
meses_portugues = {
    1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril',
    5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto',
    9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro'
}

# ===============================
# FUNÇÕES AUXILIARES
# ===============================
def extract_text_from_pdf(pdf_path):
    text = ''
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text()
    except Exception as e:
        print(f"Erro ao extrair PDF {pdf_path}: {e}")
    return text


def extract_info_from_pdf_content(pdf_content):
    match = re.search(r'Data\s+(\d{2}/\d{2}/\d{4})', pdf_content)

    if match:
        date_obj = datetime.strptime(match.group(1), '%d/%m/%Y')

        return (
            match.group(1),
            'EMAIL_REMOVIDO',
            meses_portugues[date_obj.month],
            date_obj.day,
            date_obj.year,
            f'{date_obj.day} de {meses_portugues[date_obj.month]} de {date_obj.year}'
        )

    return ('dia de mês de ano', '', 'mês', 'dia', 'ano', 'dia de mês de ano')


def number_to_currency_text_extended(number):
    inteiro = int(number)
    decimal = int(round((number - inteiro) * 100))

    texto_inteiro = num2words(inteiro, lang='pt_BR')
    texto_decimal = num2words(decimal, lang='pt_BR')

    if decimal > 0:
        return f'{texto_inteiro} reais e {texto_decimal} centavos'
    return f'{texto_inteiro} reais'


def normalize_name(name):
    name = name.lower()
    name = re.sub(r'[áàãâä]', 'a', name)
    name = re.sub(r'[éèêë]', 'e', name)
    name = re.sub(r'[íìîï]', 'i', name)
    name = re.sub(r'[óòõôö]', 'o', name)
    name = re.sub(r'[úùûü]', 'u', name)
    name = re.sub(r'[ç]', 'c', name)
    return re.sub(r'[^a-z0-9]', '', name)


def capitalize_name(name):
    return ' '.join(word.capitalize() for word in name.lower().split())

def replace_text_in_paragraph(paragraph, key, value):
    """
    Substitui uma chave (placeholder) por um valor em um parágrafo do python-docx,
    preservando formatação mesmo que o placeholder não esteja em um run específico.
    """

    if key not in paragraph.text:
        return  # nada a substituir

    # Tenta substituir apenas dentro do run que contém o placeholder
    for run in paragraph.runs:
        if key in run.text:
            run.text = run.text.replace(key, str(value))
            return  # terminou com sucesso

    # Se chegou aqui, o placeholder não estava dentro de um run específico.
    # Vamos então substituir o texto inteiro, mas mantendo a formatação base.
    new_text = paragraph.text.replace(key, str(value))

    # Salva formatação base do primeiro run (se existir)
    if paragraph.runs:
        base_run = paragraph.runs[0]
        base_font_name = base_run.font.name
        base_font_size = base_run.font.size
        base_bold = base_run.bold
        base_italic = base_run.italic
    else:
        base_font_name = None
        base_font_size = None
        base_bold = None
        base_italic = None

    # Limpa o conteúdo anterior do parágrafo
    paragraph.clear()

    # Cria novo run com o texto substituído
    new_run = paragraph.add_run(new_text)

    # Aplica a formatação base (quando existir)
    if base_font_name:
        new_run.font.name = base_font_name
    if base_font_size:
        new_run.font.size = base_font_size
    if base_bold is not None:
        new_run.bold = base_bold
    if base_italic is not None:
        new_run.italic = base_italic

def replace_paragraph_text_preserve_style(paragraph, new_text):
    """
    Substitui todo o texto do parágrafo por new_text preservando
    a formatação do primeiro run (fonte, tamanho, negrito, itálico).
    """
    # Determina formatação base (do primeiro run que tiver formatação explícita)
    base_font_name = True
    base_font_size = True
    base_bold = None
    base_italic = None

    if paragraph.runs:
        # Tenta achar um run com formatação explícita, senão usa o primeiro
        base_run = paragraph.runs[0]
        for r in paragraph.runs:
            # usa o primeiro run que tenha alguma propriedade configurada
            if (r.font.name or r.font.size or r.bold is not None or r.italic is not None):
                base_run = r
                break
        base_font_name = base_run.font.name
        base_font_size = base_run.font.size
        base_bold = base_run.bold
        base_italic = base_run.italic

    # Limpa e cria novo run com o texto substituído
    paragraph.clear()
    run = paragraph.add_run(new_text)

    # Aplica formatação base (quando disponível)
    if base_font_name:
        try:
            run.font.name = base_font_name
        except Exception:
            pass
    if base_font_size:
        try:
            run.font.size = base_font_size
        except Exception:
            pass
    # bold/italic podem ser True/False/None
    if base_bold is not None:
        run.bold = base_bold
    if base_italic is not None:
        run.italic = base_italic



# ===============================
# GERAR DOCUMENTO
# ===============================
def generate_document(row, email_date_info, current_date_info, due_date_info, template_path):

    doc = Document(template_path)

    nome = row['Funcionário']
    nome_cap = capitalize_name(nome)
    total = row['Total']

    email = row['mail'] if 'mail' in row and pd.notna(row['mail']) else ''

    _, _, mes_email, dia_email, ano_email, data_email = email_date_info
    dia_atual, mes_atual, ano_atual, data_atual = current_date_info
    dia_venc, mes_venc, ano_venc, _ = due_date_info

    replacements = {
        '[nome do servidor cap]': nome_cap,
        '[valor numérico]': f'{total:.2f}'.replace('.', ','),
        '[valor por extenso]': number_to_currency_text_extended(total),
        '[r-mail]': email,
        '[dia atual]': str(dia_atual),
        '[mês atual]': mes_atual,
        '[ano atual]': str(ano_atual),
        '[dia email]': str(dia_email),
        '[mês email]': mes_email,
        '[ano email]': str(ano_email),
    }

    for p in doc.paragraphs:
        for k, v in replacements.items():
            if k in p.text:
                p.text = p.text.replace(k, str(v))

    # Nome seguro
    safe_name = re.sub(r'[\\/*?:"<>|]', "", nome)

    output_path = os.path.join(output_dir, f'{safe_name}.docx')
    doc.save(output_path)

    print(f'✔ Gerado: {output_path}')


# ===============================
# MAIN
# ===============================
df = pd.read_excel(ods_path, engine='odf')

# Datas
today = datetime.now()
current_info = (
    today.day,
    meses_portugues[today.month],
    today.year,
    f'{today.day} de {meses_portugues[today.month]} de {today.year}'
)

last_day = calendar.monthrange(today.year, today.month)[1]
due_info = (last_day, meses_portugues[today.month], today.year, '')

# PDFs
pdf_files = [f for f in os.listdir(pdf_directory) if f.endswith('.pdf')]

pdf_map = {}

for _, row in df.iterrows():
    nome = row['Funcionário']
    nome_norm = normalize_name(nome)

    for pdf in pdf_files:
        pdf_norm = normalize_name(pdf)
        if nome_norm in pdf_norm:
            pdf_map[row['Nro Funcional']] = pdf
            break

# ===============================
# GERAÇÃO
# ===============================
for _, row in df.iterrows():

    condicao = ''
    if 'Condição' in row and pd.notna(row['Condição']):
        condicao = str(row['Condição']).strip().lower()

    nome = row['Funcionário']

    # REGRA
    if condicao in ['não enviar', 'nao enviar']:
        print(f'⛔ Pulado: {nome}')
        continue

    template = os.path.join('template', 'template_ligado.docx')
    if condicao == 'desligado':
        template = os.path.join('template', 'template_desligado.docx')

    nro = row['Nro Funcional']

    if nro in pdf_map:
        pdf_path = os.path.join(pdf_directory, pdf_map[nro])
        content = extract_text_from_pdf(pdf_path)
        email_info = extract_info_from_pdf_content(content)
    else:
        print(f'⚠ Sem PDF: {nome}')
        email_info = ('', '', 'mês', 'dia', 'ano', 'dia de mês de ano')

    generate_document(row, email_info, current_info, due_info, template)