# -*- coding: utf-8 -*-
"""
Arquivo: codigo_comentado.py
Descrição: Versão comentada do script fornecido pelo usuário. Cada bloco e função tem explicações
em português detalhando o que faz, por que está ali e observações sobre uso/possíveis pontos de
atenção / erros comuns.

Observação: este arquivo foi gerado automaticamente para fins de estudo. Não altera o comportamento
original do código — apenas adiciona comentários explicativos. Se quiser que eu simplifique,
expanda ou corrija partes do código (por exemplo melhorias de robustez), peça que eu modifique.
"""

# Importações de bibliotecas
# --------------------------
# pandas: leitura/manipulação de tabelas (Excel/ODS/CSV). Aqui usamos para abrir o arquivo .ods
# python-docx (docx): manipular documentos Word (.docx) - criar, ler e alterar parágrafos/runs
# docx.shared.Inches, Pt: unidades e tamanhos para fontes/elements no docx
# docx.enum.text.WD_ALIGN_PARAGRAPH: constantes de alinhamento (não usado no código atual, importado
#   possivelmente para alinhamentos futuros)
# re: expressões regulares para buscar padrões de texto (datas, nomes, placeholders)
# os: operações de sistema de arquivos (listar diretório, montar caminhos)
# PyPDF2: extrair texto de arquivos PDF
# num2words: converter números para texto (ex.: 123 -> "cento e vinte e três") em pt_BR
# datetime, calendar: trabalhar com datas (hoje, último dia do mês, etc.)
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
import PyPDF2
from num2words import num2words
from datetime import datetime
import calendar


# -------------------------
# MAPEAMENTO DE MESES
# -------------------------
# Dicionário usado para transformar número do mês -> nome em português (para gerar textos)
meses_portugues = {
    1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril',
    5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto',
    9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro'
}


# -------------------------
# FUNÇÕES DE EXTRAÇÃO E PROCESSAMENTO
# -------------------------

def extract_text_from_pdf(pdf_path):
    """
    Abre um PDF e concatena o texto de todas as páginas.
    Parâmetros:
        pdf_path (str): caminho para o arquivo PDF.
    Retorna:
        str: texto extraído (vazio em caso de erro).

    Observações:
    - Usa PyPDF2.PdfReader: comportamento pode variar conforme o PDF (textos extraídos
      podem ter quebras/espacamentos estranhos se o PDF vier de imagem / OCR não feito).
    - Captura exceções para evitar que um PDF com problema pare todo o processamento.
    """
    text = ''
    try:
        # abre em modo binário para leitura
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            # itera sobre todas as páginas e chama extract_text()
            for page_num in range(len(reader.pages)):
                text += reader.pages[page_num].extract_text()
    except Exception as e:
        # imprime aviso no terminal caso haja falha — útil para depuração
        print(f"Erro ao extrair texto do PDF {pdf_path}: {e}")
    return text


def extract_info_from_pdf_content(pdf_content):
    """
    Recebe o texto extraído de um PDF e tenta achar a data do email.
    Retorna uma tupla com vários campos usados no template do documento.

    Estratégia usada:
    - Busca por um padrão "Data YYYY-MM-DD" (ex.: Data 2024-07-29) usando expressões regulares.
    - Se encontrar, converte a data (via pandas.to_datetime) para extrair dia, mês e ano e traduz o mês
      para português usando meses_portugues.
    - Se não encontrar, retorna placeholders padrão (ex.: 'dia de mês de ano').

    Retorno (ordem):
      email_date (string crua extraída / placeholder),
      'EMAIL_DO_PDF_REMOVIDO' (placeholder no código original — possivelmente aqui seria o email do PDF),
      email_month_portugues (nome do mês em pt),
      email_day (dia em número),
      email_year (ano em número),
      email_date_formatted (string formatada: 'D de mês de AAAA')
    """
    # tenta encontrar 'Data YYYY-MM-DD' no conteúdo do PDF
    email_date_match = re.search(r'Data (\d{4}-\d{2}-\d{2})', pdf_content)
    
    # se houve correspondência, extrai o grupo com a data no formato ISO
    email_date = email_date_match.group(1) if email_date_match else 'dia de mês de ano'
    
    if email_date != 'dia de mês de ano':
        # converte para objeto de data (usando pandas que aceita vários formatos)
        date_obj = pd.to_datetime(email_date)
        email_month_portugues = meses_portugues[date_obj.month]
        email_day = date_obj.day
        email_year = date_obj.year
        email_date_formatted = f'{email_day} de {email_month_portugues} de {email_year}'
    else:
        # placeholders caso não tenha sido possível extrair a data
        email_month_portugues = 'mês'
        email_day = 'dia'
        email_year = 'ano'
        # email_date_formatted permanece como placeholder

    # note que a função retorna um valor fixo para o campo do email ('EMAIL_DO_PDF_REMOVIDO')
    # provavelmente o autor original quis esconder/normalizar o email extraído do PDF.
    return email_date, 'EMAIL_DO_PDF_REMOVIDO', email_month_portugues, email_day, email_year, email_date_formatted


def number_to_currency_text_extended(number):
    """
    Converte um número float (valor monetário) para a forma por extenso em português brasileiro.
    Ex.: 123.45 -> 'cento e vinte e três reais e quarenta e cinco centavos'

    Estratégia:
    - Separa parte inteira e decimal.
    - Usa num2words com lang='pt_BR' para converter as partes em palavras.
    - Combina as partes com 'reais' e 'centavos'.

    Observações/limitações:
    - Assumimos que 'number' é um float ou pode ser convertido para float.
    - O código usa round((number - inteiro) * 100) para obter os centavos: essa operação pode ter
      comportamento numérico inesperado em floats (pequenos erros de arredondamento). Em casos críticos
      é mais seguro trabalhar com inteiros de centavos (ex.: int(round(number * 100))).
    - Em caso de erro, retorna um texto de erro no formato 'VALOR_POR_EXTENSO_ERRO_xx,xx'.
    """
    try:
        inteiro = int(number)
        decimal = int(round((number - inteiro) * 100))

        texto_inteiro = num2words(inteiro, lang='pt_BR')
        texto_decimal = num2words(decimal, lang='pt_BR')

        if decimal > 0:
            return f'{texto_inteiro} reais e {texto_decimal} centavos'
        else:
            return f'{texto_inteiro} reais'
    except Exception as e:
        print(f"Erro ao converter número para extenso: {number} - {e}")
        # formata uma mensagem de erro substituindo ponto por vírgula para formato BR
        return f'VALOR_POR_EXTENSO_ERRO_{number:.2f}'.replace('.', ',')


def normalize_name_for_comparison(name):
    """
    Normaliza nomes para comparação aproximada entre strings vindas de fontes diferentes (planilha x nome do arquivo PDF).
    Operações feitas:
    - converte para minúsculas
    - remove acentos substituindo por letras ASCII correspondentes (ex.: 'ç' -> 'c')
    - remove qualquer caractere que não seja [a-z0-9] (ou seja, exclui espaços, pontuação)
    - remove espaços extras e faz strip

    Resultado: string "compacta" sem acento, usada para comparar (ex.: 'José Silva' -> 'josésilva' -> 'josesilva')

    Observações:
    - A função faz remoção manual de grupos de acentos com regex; isso funciona na maioria dos casos
      mas pode não cobrir todos os caracteres Unicode possíveis.
    - O autor chamou também re.sub(r'[^a-z0-9]', '', name) o que removeu qualquer caractere não alfanumérico
      resultando em uma string contínua sem espaços.
    """
    name = name.lower()
    name = re.sub(r'[áàãâä]', 'a', name)
    name = re.sub(r'[éèêë]', 'e', name)
    name = re.sub(r'[íìîï]', 'i', name)
    name = re.sub(r'[óòõôö]', 'o', name)
    name = re.sub(r'[úùûü]', 'u', name)
    name = re.sub(r'[ç]', 'c', name)
    # remove qualquer caractere que não seja letra minúscula ou número
    name = re.sub(r'[^a-z0-9]', '', name)
    # remove espaços repetidos e tira espaços nas extremidades (embora após a substituição anterior não haja espaços)
    name = re.sub(r'\s+', '', name).strip()
    return name


def capitalize_name(name):
    """
    Formata um nome colocando a primeira letra de cada palavra em maiúscula.
    Ex.: 'joao silva' -> 'Joao Silva'
    Observação: usa .lower() antes para garantir padronização (por exemplo nomes em MAIÚSCULAS).
    """
    return ' '.join([word.capitalize() for word in name.lower().split()])


def replace_text_in_paragraph(paragraph, key, value):
    """
    Substitui uma chave (substring) por um valor em um parágrafo do python-docx.
    Observações sobre runs (fragmentos de texto formatados):
    - Um parágrafo pode ter vários runs com formatações diferentes (negrito, itálico, fonte, etc.).
    - Se o placeholder estiver dentro de um run, a substituição atualiza apenas esse run.
    - Caso contrário, substitui o texto do parágrafo inteiro (o que pode alterar a formatação original).

    Implementação:
    - Primeiro checa se o key está no texto do parágrafo.
    - Se sim, percorre os runs procurando o run que contenha o key e substitui apenas nele.
    - Se não encontra em nenhum run, substitui no texto completo do parágrafo.
    """
    if key in paragraph.text:
        for run in paragraph.runs:
            if key in run.text:
                # substitui apenas nesse run (preserva formatação dos outros runs)
                run.text = run.text.replace(key, str(value))
                return 
        # se não encontrou run específico, substitui todo o texto do parágrafo (pode perder formatação)
        paragraph.text = paragraph.text.replace(key, str(value))


def replace_paragraph_text_preserve_style(paragraph, new_text):
    """
    Substitui TODO o texto de um parágrafo por 'new_text' preservando a formatação do primeiro run
    que tenha alguma formatação explícita (fonte, tamanho, negrito, itálico).

    Por que isso é útil?
    - python-docx não tem API direta para "mudar todo o texto mantendo estilo". Se você fizer paragraph.text = ...
      a formatação pode ser perdida. Essa função recria o parágrafo preservando atributos principais.

    Observações:
    - Tenta detectar um "base_run" com propriedades de fonte/tamanho/bold/italic; usa-o como modelo.
    - Se não houver runs, apenas cria um novo run com o texto.
    - Usa try/except para atribuir font.name/size porque nem sempre esses atributos estão disponíveis.
    """
    # Determina formatação base (do primeiro run que tiver formatação explícita)
    base_font_name = True
    base_font_size = True
    base_bold = None
    base_italic = None

    if paragraph.runs:
        # Tenta achar um run com formatação explícita, senão usa o primeiro
        base_run = paragraph.runs[0]
        for r in paragraph.runs:
            # usa o primeiro run que tenha alguma propriedade configurada
            if (r.font.name or r.font.size or r.bold is not None or r.italic is not None):
                base_run = r
                break
        base_font_name = base_run.font.name
        base_font_size = base_run.font.size
        base_bold = base_run.bold
        base_italic = base_run.italic

    # Limpa e cria novo run com o texto substituído
    paragraph.clear()
    run = paragraph.add_run(new_text)

    # Aplica formatação base (quando disponível)
    if base_font_name:
        try:
            run.font.name = base_font_name
        except Exception:
            # nem sempre é possível atribuir font.name (dependente de ambiente / docx)
            pass
    if base_font_size:
        try:
            run.font.size = base_font_size
        except Exception:
            pass
    # bold/italic podem ser True/False/None
    if base_bold is not None:
        run.bold = base_bold
    if base_italic is not None:
        run.italic = base_italic


# -------------------------
# FUNÇÃO PRINCIPAL DE GERAÇÃO DO DOCUMENTO (.docx)
# -------------------------

def generate_document(data_row, email_date_info, current_date_info, due_date_info, template_path='template.docx'):
    """
    Gera um documento Word a partir de uma linha (data_row) do DataFrame e de informações de datas.

    Parâmetros:
      data_row: Series do pandas com colunas esperadas como 'Nro Funcional', 'Funcionário', 'Total',
                'Endereço', 'Bairro', 'CEP', 'mail' etc.
      email_date_info: tupla retornada por extract_info_from_pdf_content (ou placeholder)
      current_date_info: tupla com (dia, mês_pt, ano, formatado)
      due_date_info: tupla com info do vencimento (último dia do mês, etc.)
      template_path: caminho para o arquivo .docx que serve de template

    Fluxo:
      - Abre o template com Document(template_path)
      - Extrai campos da linha do DataFrame
      - Cria um dicionário 'replacements' que mapeia placeholders para os valores reais
      - Percorre parágrafos do documento e substitui placeholders por valores
      - Salva o novo .docx com nome baseado no nome do funcionário

    Observações:
      - Caso o template contenha placeholders distribuídos em diferentes runs, replace_text_in_paragraph
        tenta preservar a formatação local.
      - Existem trechos específicos para ajustar como o parágrafo com 'Ilmo(a) Senhor(a):' é montado
        (configurações de fonte/tamanho para o nome do servidor).
      - O trecho que trata do parágrafo com texto 'Informamos que notificação semelhante...' reconstrói
        o parágrafo dividindo em partes e substituindo placeholders específicos (email e data).
    """
    # carrega o modelo
    document = Document(template_path)

    # Campos básicos da planilha
    nro_funcional = data_row['Nro Funcional']
    funcionario_raw = data_row['Funcionário']
    funcionario_capitalized = capitalize_name(funcionario_raw)
    funcionario_uppercase = funcionario_raw.upper()
    total = data_row['Total']
    
    # Campos de endereço com proteção caso a coluna não exista ou tenha NaN
    endereco_rua = data_row['Endereço'] if 'Endereço' in data_row and pd.notna(data_row['Endereço']) else ''
    bairro = data_row['Bairro'] if 'Bairro' in data_row and pd.notna(data_row['Bairro']) else ''
    complemento = data_row['complemento'] if 'complemento' in data_row and pd.notna(data_row['complemento']) else ''
    cep = data_row['CEP'] if 'CEP' in data_row and pd.notna(data_row['CEP']) else ''

    # Monta o endereço completo concatenando partes quando disponíveis
    endereco_completo = endereco_rua
    if complemento: 
        endereco_completo += f', – {complemento}'
    if bairro:
        endereco_completo += f', – {bairro}'
        # Se quiser incluir CEP, poderia concatenar aqui
        # endereco_completo += f' - CEP: {cep}'
    
    # email que virá da planilha (campo 'mail') - se não existir, usa 'mail' como placeholder
    email_address_from_excel = data_row['mail'] if 'mail' in data_row and pd.notna(data_row['mail']) else 'mail'

    # desembrulha as tuplas de data passadas para a função
    email_date_raw, _, email_month_portugues, email_day, email_year, email_date_formatted = email_date_info
    current_day, current_month_portugues, current_year, current_date_formatted = current_date_info
    due_day, due_month_portugues, due_year, due_date_formatted = due_date_info

    # dicionário com todos os placeholders que serão substituídos no documento
    replacements = {
        '[dia atual]': str(current_day),
        '[mês atual]': current_month_portugues,
        '[ano atual]': str(current_year),
        'Piracicaba, [dia atual] de [mês atual] de [ano atual].': f'Piracicaba, {current_date_formatted}.',

        '[ultimo dia do mês atual]': str(due_day),
        '[mês vencimento]': due_month_portugues,
        '[ano vencimento]': str(due_year),

        '[dia email]': str(email_day),
        '[mês email]': email_month_portugues,
        '[ano email]': str(email_year),
        '[r-mail]': email_address_from_excel,

        '[valor numérico]': f'{total:.2f}'.replace('.', ','),
        '[valor por extenso]': number_to_currency_text_extended(total),

        '[nome do servidor upper]': funcionario_uppercase,
        '[nome do servidor cap]': funcionario_capitalized,
        '[endereço do servidor]': endereco_completo,
        '[CEP do servidor]': cep,
    }

    # Percorre todos os parágrafos do documento e realiza substituições
    for paragraph in document.paragraphs:
        # Caso especial: se o parágrafo contém tanto 'Ilmo(a) Senhor(a):' quanto '[nome do servidor]'
        # o código limpa e reconstrói o parágrafo colocando o nome formatado com fonte/tamanho.
        if 'Ilmo(a) Senhor(a):' in paragraph.text and '[nome do servidor]' in paragraph.text:
            paragraph.clear()
            run_prefix = paragraph.add_run('Ilmo(a) Senhor(a):\n')
            run_prefix.bold = False
            run_name = paragraph.add_run(funcionario_capitalized)
            run_name.bold = False
            run_name.font.size = Pt(12)
            run_name.font.name = 'Calibri'
            continue

        # Caso especial: parágrafo que informa notificação por e-mail — o autor executa uma substituição
        # mais complexa porque queria inserir o e-mail e a data com certa formatação.
        if 'Informamos que notificação semelhante foi enviada ao email cadastrado no sistema ([r-mail]), em' in paragraph.text:
            # Cria uma nova lista de runs reconstruindo o parágrafo.
            new_runs = []
            temp_text = paragraph.text
            
            # coloca aqui as strings placeholders exatas esperadas no template
            email_placeholder = '[r-mail]'
            date_placeholder = 'em 20 de [mês atual] de [ano atual].'

            # divide o texto do parágrafo usando os placeholders para preservar a ordem das partes
            parts = re.split(f'({re.escape(email_placeholder)}|{re.escape(date_placeholder)})', temp_text)
            
            for part in parts:
                if part == email_placeholder:
                    new_runs.append(email_address_from_excel)
                elif part == date_placeholder:
                    new_runs.append(f'em {email_date_formatted}.')
                else:
                    new_runs.append(part)

        # Substitui todos os placeholders definidos no dicionário replacements
        for key, value in replacements.items():
            if key in paragraph.text:
                replace_text_in_paragraph(paragraph, key, value,)

    # salva com nome baseado no nome do funcionário (poderia sanitizar o nome para evitar caracteres inválidos)
    output_filename = f'{funcionario_raw.replace(" ", " ")}.docx'
    document.save(output_filename)
    print(f'Documento gerado: {output_filename}')


# -------------------------
# BLOCO PRINCIPAL (quando executado como script)
# -------------------------
if __name__ == '__main__':
    # Caminho do arquivo .ods de entrada (pode ser alterado)
    ods_path = 'teste.ods' # arquivo de entrada (excel)
    pdf_directory = '.' 
    
    # Datas atuais e informações formatadas para uso no template
    today = datetime.now()
    current_month_portugues = meses_portugues[today.month]
    current_date_formatted = f'{today.day} de {current_month_portugues} de {today.year}'
    current_date_info = (today.day, current_month_portugues, today.year, current_date_formatted)

    # calcula o último dia do mês atual (útil para vencimentos)
    last_day_of_month = calendar.monthrange(today.year, today.month)[1]
    due_date_formatted = f'{last_day_of_month} de {current_month_portugues} de {today.year}'
    due_date_info = (last_day_of_month, current_month_portugues, today.year, due_date_formatted)


    # lê a planilha ODS usando pandas com engine 'odf' (certifique-se de ter instalado 'odfpy')
    df = pd.read_excel(ods_path, engine='odf')

    # lista arquivos PDF que terminam com 'email.pdf' no diretório configurado
    pdf_files = [f for f in os.listdir(pdf_directory) if f.endswith('email.pdf')]
    
    # cria um mapeamento entre número funcional e arquivo PDF correspondente
    pdf_map = {}
    for index, row in df.iterrows():
        funcionario_nome_planilha = row['Funcionário']
        normalized_funcionario_name = normalize_name_for_comparison(funcionario_nome_planilha)
        
        print(f"\n--- Processando funcionário da planilha ---")
        print(f"Nome original da planilha: {funcionario_nome_planilha}")
        print(f"Nome normalizado da planilha: {normalized_funcionario_name}")

        found_pdf = None
        for pdf_file in pdf_files:
            # remove sufixos esperados do nome do arquivo para comparar apenas com a parte do nome
            base_pdf_name = pdf_file.replace('_email.pdf', '').replace('_TANCREDO', '')
            normalized_pdf_filename = normalize_name_for_comparison(base_pdf_name)
            
            print(f"  Comparando com PDF: {pdf_file}")
            print(f"  Nome original do PDF (base): {base_pdf_name}")
            print(f"  Nome normalizado do PDF: {normalized_pdf_filename}")
            
            # estratégia de comparação: checar se o nome normalizado da planilha está dentro do nome do PDF
            if normalized_funcionario_name in normalized_pdf_filename or normalized_pdf_filename.startswith(normalized_funcionario_name):
                found_pdf = pdf_file
                print(f"  *** Correspondência encontrada: {pdf_file} ***")
                break
        
        # se encontrou, registra no mapa (chave: Nro Funcional)
        if found_pdf:
            pdf_map[row['Nro Funcional']] = found_pdf
        else:
            print(f"Aviso: Nenhum PDF de email correspondente encontrado para o funcionário: {funcionario_nome_planilha} (Nro Funcional: {row['Nro Funcional']})")

    # Agora, para cada linha do DataFrame, gera um documento .docx substituindo as infos
    print(f"\n--- Iniciando geração de documentos ---")
    for index, row in df.iterrows():
        nro_funcional = row['Nro Funcional']
        email_address_from_excel = row['mail'] if 'mail' in row and pd.notna(row['mail']) else 'r-mail'

        if nro_funcional in pdf_map:
            # obtém o caminho para o PDF correspondente e extrai seu conteúdo
            current_pdf_path = os.path.join(pdf_directory, pdf_map[nro_funcional])
            pdf_content = extract_text_from_pdf(current_pdf_path)
            # extrai a data do email a partir do conteúdo do PDF
            email_date_raw, _, email_month_portugues, email_day, email_year, email_date_formatted = extract_info_from_pdf_content(pdf_content)
            email_date_info = (email_date_raw, email_address_from_excel, email_month_portugues, email_day, email_year, email_date_formatted)
            # gera o documento usando o template
            generate_document(row, email_date_info, current_date_info, due_date_info, template_path='template.docx')
        else:
            # Se não houver PDF, ainda podemos gerar o documento com placeholders padrão
            print(f"Aviso: Nenhum PDF de email encontrado para o funcionário: {row['Funcionário']} (Nro Funcional: {nro_funcional}). Gerando documento com data de email padrão.")
            email_date_info = ('dia de mês de ano', email_address_from_excel, 'mês', 'dia', 'ano', 'dia de mês de ano')
            generate_document(row, email_date_info, current_date_info, due_date_info, template_path='template.docx')
