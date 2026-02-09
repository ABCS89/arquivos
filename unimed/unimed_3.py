import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
import PyPDF2
from num2words import num2words
from datetime import datetime
import calendar

# Mapeamento de meses para português
meses_portugues = {
    1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril',
    5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto',
    9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro'
}

def extract_text_from_pdf(pdf_path):
    text = ''
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                text += reader.pages[page_num].extract_text()
    except Exception as e:
        print(f"Erro ao extrair texto do PDF {pdf_path}: {e}")
    return text

def extract_info_from_pdf_content(pdf_content):
    # Procura pela data no formato AAAA-MM-DD
    email_date_match = re.search(r'Data (\d{4}-\d{2}-\d{2})', pdf_content)
    
    # REGEX para capturar todos os e-mails
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    todos_emails = re.findall(email_pattern, pdf_content)
    
    # Filtra para ignorar o e-mail do remetente (DRH)
    emails_filtrados = [e for e in todos_emails if 'drh.pagamento' not in e.lower()]
    
    # Define o e-mail do servidor (o primeiro que sobrar após o filtro)
    email_pdf_extraido = emails_filtrados[0] if emails_filtrados else None
    
    email_date = email_date_match.group(1) if email_date_match else 'dia de mês de ano'
    
    if email_date != 'dia de mês de ano':
        date_obj = pd.to_datetime(email_date)
        email_month_portugues = meses_portugues[date_obj.month]
        email_day = date_obj.day
        email_year = date_obj.year
        email_date_formatted = f'{email_day} de {email_month_portugues} de {email_year}'
    else:
        email_month_portugues = 'mês'
        email_day = 'dia'
        email_year = 'ano'
        email_date_formatted = 'dia de mês de ano'

    return email_date, 'EMAIL_DO_PDF_REMOVIDO', email_month_portugues, email_day, email_year, email_date_formatted

def number_to_currency_text_extended(number):
    try:
        inteiro = int(number)
        decimal = int(round((number - inteiro) * 100))

        texto_inteiro = num2words(inteiro, lang='pt_BR')
        texto_decimal = num2words(decimal, lang='pt_BR')

        if decimal > 0:
            return f'{texto_inteiro} reais e {texto_decimal} centavos'
        else:
            return f'{texto_inteiro} reais'
    except Exception as e:
        print(f"Erro ao converter número para extenso: {number} - {e}")
        return f'VALOR_POR_EXTENSO_ERRO_{number:.2f}'.replace('.', ',')

def normalize_name_for_comparison(name):
    name = name.lower()
    name = re.sub(r'[áàãâä]', 'a', name)
    name = re.sub(r'[éèêë]', 'e', name)
    name = re.sub(r'[íìîï]', 'i', name)
    name = re.sub(r'[óòõôö]', 'o', name)
    name = re.sub(r'[úùûü]', 'u', name)
    name = re.sub(r'[ç]', 'c', name)
    name = re.sub(r'[^a-z0-9]', '', name) 
    name = re.sub(r'\s+', '', name).strip()
    return name

def capitalize_name(name):
    return ' '.join([word.capitalize() for word in name.lower().split()])

def replace_text_in_paragraph(paragraph, key, value):
    """
    Substitui uma chave (placeholder) por um valor em um parágrafo do python-docx,
    preservando formatação mesmo que o placeholder não esteja em um run específico.
    """

    if key not in paragraph.text:
        return  # nada a substituir

    # Tenta substituir apenas dentro do run que contém o placeholder
    for run in paragraph.runs:
        if key in run.text:
            run.text = run.text.replace(key, str(value))
            return  # terminou com sucesso

    # Se chegou aqui, o placeholder não estava dentro de um run específico.
    # Vamos então substituir o texto inteiro, mas mantendo a formatação base.
    new_text = paragraph.text.replace(key, str(value))

    # Salva formatação base do primeiro run (se existir)
    if paragraph.runs:
        base_run = paragraph.runs[0]
        base_font_name = base_run.font.name
        base_font_size = base_run.font.size
        base_bold = base_run.bold
        base_italic = base_run.italic
    else:
        base_font_name = None
        base_font_size = None
        base_bold = None
        base_italic = None

    # Limpa o conteúdo anterior do parágrafo
    paragraph.clear()

    # Cria novo run com o texto substituído
    new_run = paragraph.add_run(new_text)

    # Aplica a formatação base (quando existir)
    if base_font_name:
        new_run.font.name = base_font_name
    if base_font_size:
        new_run.font.size = base_font_size
    if base_bold is not None:
        new_run.bold = base_bold
    if base_italic is not None:
        new_run.italic = base_italic

def replace_paragraph_text_preserve_style(paragraph, new_text):
    """
    Substitui todo o texto do parágrafo por new_text preservando
    a formatação do primeiro run (fonte, tamanho, negrito, itálico).
    """
    # Determina formatação base (do primeiro run que tiver formatação explícita)
    base_font_name = True
    base_font_size = True
    base_bold = None
    base_italic = None

    if paragraph.runs:
        # Tenta achar um run com formatação explícita, senão usa o primeiro
        base_run = paragraph.runs[0]
        for r in paragraph.runs:
            # usa o primeiro run que tenha alguma propriedade configurada
            if (r.font.name or r.font.size or r.bold is not None or r.italic is not None):
                base_run = r
                break
        base_font_name = base_run.font.name
        base_font_size = base_run.font.size
        base_bold = base_run.bold
        base_italic = base_run.italic

    # Limpa e cria novo run com o texto substituído
    paragraph.clear()
    run = paragraph.add_run(new_text)

    # Aplica formatação base (quando disponível)
    if base_font_name:
        try:
            run.font.name = base_font_name
        except Exception:
            pass
    if base_font_size:
        try:
            run.font.size = base_font_size
        except Exception:
            pass
    # bold/italic podem ser True/False/None
    if base_bold is not None:
        run.bold = base_bold
    if base_italic is not None:
        run.italic = base_italic




def generate_document(data_row, email_date_info, current_date_info, due_date_info, template_path='template.docx'):
    document = Document(template_path)

    nro_funcional = data_row['Nro Funcional']
    funcionario_raw = data_row['Funcionário']
    funcionario_capitalized = capitalize_name(funcionario_raw)
    funcionario_uppercase = funcionario_raw.upper()
    total = data_row['Total']
    
    endereco_rua = data_row['Endereço'] if 'Endereço' in data_row and pd.notna(data_row['Endereço']) else ''
    bairro = data_row['Bairro'] if 'Bairro' in data_row and pd.notna(data_row['Bairro']) else ''
    complemento = data_row['Complemento'] if 'Complemento' in data_row and pd.notna(data_row['Complemento']) else ''
    cep = data_row['CEP'] if 'CEP' in data_row and pd.notna(data_row['CEP']) else ''

    endereco_completo = endereco_rua
    if complemento: 
        endereco_completo += f' – {complemento}'
    if bairro:
        endereco_completo += f' – {bairro}'
        # endereco_completo += f' - CEP: {cep}'
    
    email_address_from_excel = data_row['mail'] if 'mail' in data_row and pd.notna(data_row['mail']) else 'mail'

    email_date_raw, _, email_month_portugues, email_day, email_year, email_date_formatted = email_date_info
    current_day, current_month_portugues, current_year, current_date_formatted = current_date_info
    due_day, due_month_portugues, due_year, due_date_formatted = due_date_info

    
    replacements = {
        '[dia atual]': str(current_day),
        '[mês atual]': current_month_portugues,
        '[ano atual]': str(current_year),
        'Piracicaba, [dia atual] de [mês atual] de [ano atual].': f'Piracicaba, {current_date_formatted}.',

        '[ultimo dia do mês atual]': str(due_day),
        '[mês vencimento]': due_month_portugues,
        '[ano vencimento]': str(due_year),

        '[dia email]': str(email_day),
        '[mês email]': email_month_portugues,
        '[ano email]': str(email_year),
        '[r-mail]': email_address_from_excel,

        '[valor numérico]': f'{total:.2f}'.replace('.', ','),
        '[valor por extenso]': number_to_currency_text_extended(total),

        '[nome do servidor upper]': funcionario_uppercase,
        '[nome do servidor cap]': funcionario_capitalized,
        '[endereço do servidor]': endereco_completo,
        '[CEP do servidor]': cep,
    }

    for paragraph in document.paragraphs:
        if 'Ilmo(a) Senhor(a):' in paragraph.text and '[nome do servidor]' in paragraph.text:
            paragraph.clear()
            run_prefix = paragraph.add_run('Ilmo(a) Senhor(a):\n')
            run_prefix.bold = False
            run_name = paragraph.add_run(funcionario_capitalized)
            run_name.bold = False
            run_name.font.size = Pt(12)
            run_name.font.name = 'Calibri'
            continue

        # 2. Correção de E-mail e Data da Notificação Anterior
        if 'Informamos que notificação semelhante foi enviada ao email cadastrado no sistema ([r-mail]), em' in paragraph.text:
            # Criar uma nova lista de runs para reconstruir o parágrafo
            new_runs = []
            temp_text = paragraph.text
            
            
            # Encontrar a posição do placeholder do email e da data
            email_placeholder = '[r-mail]'
            date_placeholder = 'em 20 de [mês atual] de [ano atual].'

            # Dividir o texto do parágrafo em partes antes, durante e depois dos placeholders
            parts = re.split(f'({re.escape(email_placeholder)}|{re.escape(date_placeholder)})', temp_text)
            
            for part in parts:
                if part == email_placeholder:
                    new_runs.append(email_address_from_excel)
                elif part == date_placeholder:
                    new_runs.append(f'em {email_date_formatted}.')
                else:
                    new_runs.append(part)        # 3. Substituições Gerais
        for key, value in replacements.items():
            replace_text_in_paragraph(paragraph, key, value)

    output_filename = f'{funcionario_raw}.docx'
    document.save(output_filename)
    print(f'Sucesso: {output_filename}')

if __name__ == '__main__':
    ods_path = 'teste.ods' # arquivo de entrada (excel)
    pdf_directory = '.' 
    
    today = datetime.now()
    current_month_portugues = meses_portugues[today.month]
    current_date_formatted = f'{today.day} de {current_month_portugues} de {today.year}'
    current_date_info = (today.day, current_month_portugues, today.year, current_date_formatted)

    last_day_of_month = calendar.monthrange(today.year, today.month)[1]
    due_date_formatted = f'{last_day_of_month} de {current_month_portugues} de {today.year}'
    due_date_info = (last_day_of_month, current_month_portugues, today.year, due_date_formatted)


    df = pd.read_excel(ods_path, engine='odf')

    pdf_files = [f for f in os.listdir(pdf_directory) if f.endswith('email.pdf')]
    
    pdf_files = [f for f in os.listdir('.') if f.endswith('email.pdf')]
    
    # Mapeamento
    pdf_map = {}
    for _, row in df.iterrows():
        n_planilha = normalize_name_for_comparison(str(row['Funcionário']))
        for pf in pdf_files:
            n_pdf = normalize_name_for_comparison(pf.replace('_email.pdf', '').replace('_TANCREDO', ''))
            if n_planilha in n_pdf or n_pdf.startswith(n_planilha):
                pdf_map[row['Nro Funcional']] = pf
                break

    for _, row in df.iterrows():
        nro = row['Nro Funcional']
        if nro in pdf_map:
            content = extract_text_from_pdf(pdf_map[nro])
            e_raw, e_mail, e_m, e_d, e_y, e_fmt = extract_info_from_pdf_content(content)
            email_info = (e_raw, e_mail, e_m, e_d, e_y, e_fmt)
        else:
            email_info = ('dia de mês de ano', None, 'mês', 'dia', 'ano', 'dia de mês de ano')
        
        generate_document(row, email_info, cur_info, due_info)