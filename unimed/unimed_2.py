import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
import PyPDF2
from num2words import num2words
from datetime import datetime
import calendar

# Mapeamento de meses para português
meses_portugues = {
    1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril',
    5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto',
    9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro'
}

def extract_text_from_pdf(pdf_path):
    text = ''
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                text += reader.pages[page_num].extract_text()
    except Exception as e:
        print(f"Erro ao extrair texto do PDF {pdf_path}: {e}")
    return text

def extract_info_from_pdf_content(pdf_content):
    # Procura pela data
    email_date_match = re.search(r'Data (\d{4}-\d{2}-\d{2})', pdf_content)
    
    # --- ALTERAÇÃO AQUI: Busca o e-mail real no conteúdo do PDF ---
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    email_match = re.search(email_pattern, pdf_content)
    email_pdf_extraido = email_match.group(0) if email_match else None
    
    email_date = email_date_match.group(1) if email_date_match else 'dia de mês de ano'
    
    if email_date != 'dia de mês de ano':
        date_obj = pd.to_datetime(email_date)
        email_month_portugues = meses_portugues[date_obj.month]
        email_day = date_obj.day
        email_year = date_obj.year
        email_date_formatted = f'{email_day} de {email_month_portugues} de {email_year}'
    else:
        email_month_portugues = 'mês'
        email_day = 'dia'
        email_year = 'ano'
        email_date_formatted = 'dia de mês de ano'

    # Retorna o e-mail extraído em vez de texto fixo
    return email_date, email_pdf_extraido, email_month_portugues, email_day, email_year, email_date_formatted

def number_to_currency_text_extended(number):
    try:
        inteiro = int(number)
        decimal = int(round((number - inteiro) * 100))
        texto_inteiro = num2words(inteiro, lang='pt_BR')
        texto_decimal = num2words(decimal, lang='pt_BR')
        if decimal > 0:
            return f'{texto_inteiro} reais e {texto_decimal} centavos'
        else:
            return f'{texto_inteiro} reais'
    except Exception as e:
        print(f"Erro ao converter número para extenso: {number} - {e}")
        return f'VALOR_POR_EXTENSO_ERRO_{number:.2f}'.replace('.', ',')

def normalize_name_for_comparison(name):
    name = name.lower()
    name = re.sub(r'[áàãâä]', 'a', name)
    name = re.sub(r'[éèêë]', 'e', name)
    name = re.sub(r'[íìîï]', 'i', name)
    name = re.sub(r'[óòõôö]', 'o', name)
    name = re.sub(r'[úùûü]', 'u', name)
    name = re.sub(r'[ç]', 'c', name)
    name = re.sub(r'[^a-z0-9]', '', name) 
    name = re.sub(r'\s+', '', name).strip()
    return name

def capitalize_name(name):
    return ' '.join([word.capitalize() for word in name.lower().split()])

def replace_text_in_paragraph(paragraph, key, value):
    if key not in paragraph.text:
        return
    for run in paragraph.runs:
        if key in run.text:
            run.text = run.text.replace(key, str(value))
            return
    new_text = paragraph.text.replace(key, str(value))
    if paragraph.runs:
        base_run = paragraph.runs[0]
        base_font_name = base_run.font.name
        base_font_size = base_run.font.size
        base_bold = base_run.bold
        base_italic = base_run.italic
    else:
        base_font_name = base_font_size = base_bold = base_italic = None

    paragraph.clear()
    new_run = paragraph.add_run(new_text)
    if base_font_name: new_run.font.name = base_font_name
    if base_font_size: new_run.font.size = base_font_size
    if base_bold is not None: new_run.bold = base_bold
    if base_italic is not None: new_run.italic = base_italic

def generate_document(data_row, email_date_info, current_date_info, due_date_info, template_path='template.docx'):
    document = Document(template_path)

    # --- ALTERAÇÃO AQUI: Desempacota o email vindo do PDF ---
    email_date_raw, email_do_pdf, email_month_portugues, email_day, email_year, email_date_formatted = email_date_info

    # Lógica de prioridade: PDF primeiro, planilha depois
    email_address_from_excel = data_row['mail'] if 'mail' in data_row and pd.notna(data_row['mail']) else 'mail'
    email_final = email_do_pdf if email_do_pdf else email_address_from_excel

    funcionario_raw = data_row['Funcionário']
    funcionario_capitalized = capitalize_name(funcionario_raw)
    funcionario_uppercase = funcionario_raw.upper()
    total = data_row['Total']
    
    endereco_rua = data_row['Endereço'] if 'Endereço' in data_row and pd.notna(data_row['Endereço']) else ''
    bairro = data_row['Bairro'] if 'Bairro' in data_row and pd.notna(data_row['Bairro']) else ''
    complemento = data_row['Complemento'] if 'Complemento' in data_row and pd.notna(data_row['Complemento']) else ''
    cep = data_row['CEP'] if 'CEP' in data_row and pd.notna(data_row['CEP']) else ''

    endereco_completo = endereco_rua
    if complemento: endereco_completo += f' – {complemento}'
    if bairro: endereco_completo += f' – {bairro}'
    
    current_day, current_month_portugues, current_year, current_date_formatted = current_date_info
    due_day, due_month_portugues, due_year, due_date_formatted = due_date_info

    replacements = {
        '[dia atual]': str(current_day),
        '[mês atual]': current_month_portugues,
        '[ano atual]': str(current_year),
        'Piracicaba, [dia atual] de [mês atual] de [ano atual].': f'Piracicaba, {current_date_formatted}.',
        '[ultimo dia do mês atual]': str(due_day),
        '[mês vencimento]': due_month_portugues,
        '[ano vencimento]': str(due_year),
        '[dia email]': str(email_day),
        '[mês email]': email_month_portugues,
        '[ano email]': str(email_year),
        '[r-mail]': email_final, # USA O EMAIL FINAL (PDF OU EXCEL)
        '[valor numérico]': f'{total:.2f}'.replace('.', ','),
        '[valor por extenso]': number_to_currency_text_extended(total),
        '[nome do servidor upper]': funcionario_uppercase,
        '[nome do servidor cap]': funcionario_capitalized,
        '[endereço do servidor]': endereco_completo,
        '[CEP do servidor]': cep,
    }

    for paragraph in document.paragraphs:
        if 'Ilmo(a) Senhor(a):' in paragraph.text and '[nome do servidor]' in paragraph.text:
            paragraph.clear()
            paragraph.add_run('Ilmo(a) Senhor(a):\n').bold = False
            run_name = paragraph.add_run(funcionario_capitalized)
            run_name.bold = False
            run_name.font.size = Pt(12)
            run_name.font.name = 'Calibri'
            continue

        if 'Informamos que notificação semelhante foi enviada ao email cadastrado no sistema ([r-mail]), em' in paragraph.text:
            # Substituição direta no texto para o parágrafo especial
            paragraph.text = paragraph.text.replace('[r-mail]', email_final)
            paragraph.text = paragraph.text.replace('em 20 de [mês atual] de [ano atual].', f'em {email_date_formatted}.')

        for key, value in replacements.items():
            if key in paragraph.text:
                replace_text_in_paragraph(paragraph, key, value)

    output_filename = f'{funcionario_raw}.docx'
    document.save(output_filename)
    print(f'Documento gerado: {output_filename}')

if __name__ == '__main__':
    ods_path = 'teste.ods' 
    pdf_directory = '.' 
    
    today = datetime.now()
    current_month_portugues = meses_portugues[today.month]
    current_date_formatted = f'{today.day} de {current_month_portugues} de {today.year}'
    current_date_info = (today.day, current_month_portugues, today.year, current_date_formatted)

    last_day_of_month = calendar.monthrange(today.year, today.month)[1]
    due_date_formatted = f'{last_day_of_month} de {current_month_portugues} de {today.year}'
    due_date_info = (last_day_of_month, current_month_portugues, today.year, due_date_formatted)

    df = pd.read_excel(ods_path, engine='odf')
    pdf_files = [f for f in os.listdir(pdf_directory) if f.endswith('email.pdf')]
    
    pdf_map = {}
    for index, row in df.iterrows():
        norm_planilha = normalize_name_for_comparison(row['Funcionário'])
        for pdf_file in pdf_files:
            norm_pdf = normalize_name_for_comparison(pdf_file.replace('_email.pdf', '').replace('_TANCREDO', ''))
            if norm_planilha in norm_pdf or norm_pdf.startswith(norm_planilha):
                pdf_map[row['Nro Funcional']] = pdf_file
                break
    
    for index, row in df.iterrows():
        nro_funcional = row['Nro Funcional']
        if nro_funcional in pdf_map:
            current_pdf_path = os.path.join(pdf_directory, pdf_map[nro_funcional])
            pdf_content = extract_text_from_pdf(current_pdf_path)
            
            # --- ALTERAÇÃO AQUI: Captura o email extraído do PDF ---
            e_date_raw, e_mail_pdf, e_month, e_day, e_year, e_fmt = extract_info_from_pdf_content(pdf_content)
            email_date_info = (e_date_raw, e_mail_pdf, e_month, e_day, e_year, e_fmt)
            
            generate_document(row, email_date_info, current_date_info, due_date_info, template_path='template.docx')
        else:
            email_excel = row['mail'] if 'mail' in row and pd.notna(row['mail']) else 'mail'
            email_date_info = ('dia de mês de ano', None, 'mês', 'dia', 'ano', 'dia de mês de ano')
            generate_document(row, email_date_info, current_date_info, due_date_info, template_path='template.docx')