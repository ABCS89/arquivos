import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
import PyPDF2
from num2words import num2words
from datetime import datetime
import calendar

# Mapeamento de meses para português
meses_portugues = {
    1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril',
    5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto',
    9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro'
}

def extract_text_from_pdf(pdf_path):
    text = ''
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                text += reader.pages[page_num].extract_text()
    except Exception as e:
        print(f"Erro ao extrair texto do PDF {pdf_path}: {e}")
    return text

def extract_info_from_pdf_content(pdf_content):
    email_date_match = re.search(r'Data (\d{4}-\d{2}-\d{2})', pdf_content)
    email_address_match = re.search(r'Para <([^>]+)>', pdf_content)
    
    email_date = email_date_match.group(1) if email_date_match else 'DATA_NAO_ENCONTRADA'
    email_address = email_address_match.group(1) if email_address_match else 'EMAIL_NAO_ENCONTRADO'
    
    if email_date != 'DATA_NAO_ENCONTRADA':
        date_obj = pd.to_datetime(email_date)
        email_month_portugues = meses_portugues[date_obj.month]
        email_day = date_obj.day
        email_year = date_obj.year
        email_date_formatted = f'{email_day} de {email_month_portugues} de {email_year}'
    else:
        email_month_portugues = 'MES_NAO_ENCONTRADO'
        email_day = 'DIA_NAO_ENCONTRADO'
        email_year = 'ANO_NAO_ENCONTRADO'
        email_date_formatted = 'DATA_NAO_ENCONTRADA'

    return email_date, email_address, email_month_portugues, email_day, email_year, email_date_formatted

def number_to_currency_text_extended(number):
    try:
        inteiro = int(number)
        decimal = int(round((number - inteiro) * 100))

        texto_inteiro = num2words(inteiro, lang='pt_BR')
        texto_decimal = num2words(decimal, lang='pt_BR')

        if decimal > 0:
            return f'{texto_inteiro} reais e {texto_decimal} centavos'
        else:
            return f'{texto_inteiro} reais'
    except Exception as e:
        print(f"Erro ao converter número para extenso: {number} - {e}")
        return f'VALOR_POR_EXTENSO_ERRO_{number:.2f}'.replace('.', ',')

def normalize_name(name):
    # Remove acentos, caracteres especiais, múltiplos espaços e converte para minúsculas
    name = name.lower()
    name = re.sub(r'[áàãâä]', 'a', name)
    name = re.sub(r'[éèêë]', 'e', name)
    name = re.sub(r'[íìîï]', 'i', name)
    name = re.sub(r'[óòõôö]', 'o', name)
    name = re.sub(r'[úùûü]', 'u', name)
    name = re.sub(r'[ç]', 'c', name)
    name = re.sub(r'[^a-z0-9]', '', name) # Remove qualquer coisa que não seja letra ou número
    name = re.sub(r'\s+', '', name).strip() # Remove todos os espaços
    return name

def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for i in range(len(inline)):
            if key in inline[i].text:
                text = inline[i].text.replace(key, str(value))
                inline[i].text = text
                return 
        paragraph.text = paragraph.text.replace(key, str(value))

def generate_document(data_row, email_info, current_date_info, due_date_info, template_path='template.docx'):
    document = Document(template_path)

    nro_funcional = data_row['Nro Funcional']
    funcionario = data_row['Funcionário']
    total = data_row['Total']
    
    endereco = data_row['Endereço'] if 'Endereço' in data_row and pd.notna(data_row['Endereço']) else 'ENDERECO_NAO_ENCONTRADO'
    cep = data_row['CEP'] if 'CEP' in data_row and pd.notna(data_row['CEP']) else 'CEP_NAO_ENCONTRADO'

    email_date_raw, email_address, email_month_portugues, email_day, email_year, email_date_formatted = email_info
    current_day, current_month_portugues, current_year, current_date_formatted = current_date_info
    due_day, due_month_portugues, due_year, due_date_formatted = due_date_info

    replacements = {
        '[dia atual]': str(current_day),
        '[mês atual]': current_month_portugues,
        '[ano atual]': str(current_year),
        'Piracicaba, [dia atual] de [mês atual] de [ano atual].': f'Piracicaba, {current_date_formatted}.',

        '[ultimo dia do mês atual]': str(due_day),
        '[mês vencimento]': due_month_portugues,
        '[ano vencimento]': str(due_year),

        '[dia email]': str(email_day),
        '[mês email]': email_month_portugues,
        '[ano email]': str(email_year),
        '[endereço de e-mail]': email_address,

        '[valor numérico]': f'{total:.2f}'.replace('.', ','),
        '[valor por extenso]': number_to_currency_text_extended(total),

        '[nome do servidor]': funcionario,
        '[endereço do servidor]': endereco,
        '[CEP do servidor]': cep,
        'V. Sª.': f'V. Sa. {funcionario}',
    }

    for paragraph in document.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                if key == '[nome do servidor]' and paragraph.text.strip().startswith('Ilmo(a) Senhor(a):'):
                    paragraph.clear()
                    paragraph.add_run('Ilmo(a) Senhor(a):\n').bold = False
                    run = paragraph.add_run(str(value))
                    run.bold = True
                elif key == 'Piracicaba, [dia atual] de [mês atual] de [ano atual].':
                    # Special handling for the header date to replace the full string
                    paragraph.text = paragraph.text.replace(key, str(value))
                else:
                    replace_text_in_paragraph(paragraph, key, value)
        
        if '[endereço do servidor]' in paragraph.text and endereco == 'ENDERECO_NAO_ENCONTRADO':
            replace_text_in_paragraph(paragraph, '[endereço do servidor]', '')
        if '[CEP do servidor]' in paragraph.text and cep == 'CEP_NAO_ENCONTRADO':
            replace_text_in_paragraph(paragraph, '[CEP do servidor]', '')

    output_filename = f'documento_{nro_funcional}_{funcionario.replace(" ", "_")}.docx'
    document.save(output_filename)
    print(f'Documento gerado: {output_filename}')

if __name__ == '__main__':
    ods_path = 'teste.ods'
    pdf_directory = '.' 
    
    today = datetime.now()
    current_month_portugues = meses_portugues[today.month]
    current_date_formatted = f'{today.day} de {current_month_portugues} de {today.year}'
    current_date_info = (today.day, current_month_portugues, today.year, current_date_formatted)

    last_day_of_month = calendar.monthrange(today.year, today.month)[1]
    due_date_formatted = f'{last_day_of_month} de {current_month_portugues} de {today.year}'
    due_date_info = (last_day_of_month, current_month_portugues, today.year, due_date_formatted)

    if not os.path.exists('template.docx'):
        doc_template = Document()
        doc_template.add_paragraph('Prefeitura do Município de Piracicaba')
        doc_template.add_paragraph('SECRETARIA DE ADMINISTRAÇÃO E GOVERNO')
        doc_template.add_paragraph('Gerência de Recursos Humanos -')
        doc_template.add_paragraph('')
        doc_template.add_paragraph('Piracicaba, [dia atual] de [mês atual] de [ano atual].') # Updated header date
        doc_template.add_paragraph('')
        doc_template.add_paragraph('Assunto: pagamento do Plano de Saúde dos Servidores Públicos de Piracicaba')
        doc_template.add_paragraph('')
        doc_template.add_paragraph('Prezado(a) Senhor(a):')
        doc_template.add_paragraph('')
        doc_template.add_paragraph('Pelo presente, solicitamos o pagamento do boleto em anexo, com vencimento em [ultimo dia do mês atual] de [mês vencimento] de [ano vencimento], referente às mensalidades e coparticipações do Plano de Saúde dos Servidores Públicos de Piracicaba, no valor de R$ [valor numérico] ([valor por extenso]).') # Updated due date
        doc_template.add_paragraph('Informamos que notificação semelhante foi enviada ao e-mail cadastrado no sistema ([endereço de e-mail]), em [dia email] de [mês email] de [ano email].') # Email date
        doc_template.add_paragraph('Ressaltamos que o não pagamento poderá implicar na rescisão do plano de saúde, conforme dispositivos legais vigentes.')
        doc_template.add_paragraph('Aproveitamos a oportunidade para renovar a V. Sª., os protestos de consideração.')
        doc_template.add_paragraph('')
        doc_template.add_paragraph('Atenciosamente,')
        doc_template.add_paragraph('')
        doc_template.add_paragraph('JOSIEL WILLIAM PAES RODRIGUES')
        doc_template.add_paragraph('Chefe de Setor')
        doc_template.add_paragraph('')
        doc_template.add_paragraph('VISTO')
        doc_template.add_paragraph('')
        doc_template.add_paragraph('PAULO SERGIO MILANEZ FILHO')
        doc_template.add_paragraph('Gestor de Unidade')
        doc_template.add_paragraph('')
        doc_template.add_paragraph('Ilmo(a) Senhor(a):')
        doc_template.add_paragraph('[nome do servidor]')
        doc_template.add_paragraph('Rua: [endereço do servidor]')
        doc_template.add_paragraph('CEP: [CEP do servidor]')
        doc_template.save('template.docx')

    df = pd.read_excel(ods_path, engine='odf')

    pdf_files = [f for f in os.listdir(pdf_directory) if f.endswith('email.pdf')]
    
    pdf_map = {}
    for index, row in df.iterrows():
        funcionario_nome_planilha = row['Funcionário']
        normalized_funcionario_name = normalize_name(funcionario_nome_planilha)
        
        print(f"\n--- Processando funcionário da planilha ---")
        print(f"Nome original da planilha: {funcionario_nome_planilha}")
        print(f"Nome normalizado da planilha: {normalized_funcionario_name}")

        found_pdf = None
        for pdf_file in pdf_files:
            base_pdf_name = pdf_file.replace('_email.pdf', '').replace('_TANCREDO', '')
            normalized_pdf_filename = normalize_name(base_pdf_name)
            
            print(f"  Comparando com PDF: {pdf_file}")
            print(f"  Nome original do PDF (base): {base_pdf_name}")
            print(f"  Nome normalizado do PDF: {normalized_pdf_filename}")
            
            # Lógica de correspondência aprimorada: verificar se o nome normalizado da planilha está contido no nome normalizado do PDF
            # ou se o nome normalizado do PDF começa com o nome normalizado do funcionário (para casos onde o nome do PDF é mais curto)
            if normalized_funcionario_name in normalized_pdf_filename or normalized_pdf_filename.startswith(normalized_funcionario_name):
                found_pdf = pdf_file
                print(f"  *** Correspondência encontrada: {pdf_file} ***")
                break
        
        if found_pdf:
            pdf_map[row['Nro Funcional']] = found_pdf
        else:
            print(f"Aviso: Nenhum PDF de e-mail correspondente encontrado para o funcionário: {funcionario_nome_planilha} (Nro Funcional: {row['Nro Funcional']})")

    print(f"\n--- Iniciando geração de documentos ---")
    for index, row in df.iterrows():
        nro_funcional = row['Nro Funcional']
        if nro_funcional in pdf_map:
            current_pdf_path = os.path.join(pdf_directory, pdf_map[nro_funcional])
            pdf_content = extract_text_from_pdf(current_pdf_path)
            email_date_raw, email_address, email_month_portugues, email_day, email_year, email_date_formatted = extract_info_from_pdf_content(pdf_content)
            email_info = (email_date_raw, email_address, email_month_portugues, email_day, email_year, email_date_formatted)
            generate_document(row, email_info, current_date_info, due_date_info, template_path='template.docx')
        else:
            print(f"Aviso: Nenhum PDF de e-mail encontrado para o funcionário: {row['Funcionário']} (Nro Funcional: {nro_funcional}). Documento para este funcionário não será gerado.")

