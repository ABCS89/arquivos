import pandas as pd
from docx import Document
import os
from datetime import datetime
from num2words import num2words

# ========= CONFIGURAÇÕES ========= #

arquivo_excel = os.path.join('../template/devedores.xlsx')
template_path = os.path.join('../template/template_refis.docx')
pasta_saida = os.path.join('../output/refis')

os.makedirs(pasta_saida, exist_ok=True)

# ========= FUNÇÕES ========= #

def normalize_name(nome):
    import re
    nome = str(nome).upper()
    nome = re.sub(r'[^A-Z0-9 ]', '', nome)
    nome = re.sub(r'\s+', '_', nome)
    return nome.strip('_')

def formatar_moeda(valor):
    try:
        valor = float(valor)
        return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except:
        return "R$ 0,00"

def substituir_texto(doc, substituicoes):
    for p in doc.paragraphs:
        for key, value in substituicoes.items():
            if key in p.text:
                for run in p.runs:
                    run.text = run.text.replace(key, str(value))

    # também substitui dentro de tabelas (caso use no template)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in substituicoes.items():
                    if key in cell.text:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.text = run.text.replace(key, str(value))

# ========= DATA ATUAL ========= #

hoje = datetime.now()

meses = {
    1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril',
    5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto',
    9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro'
}

dia = hoje.day
mes = meses[hoje.month]
ano = hoje.year

# ========= LEITURA DA PLANILHA ========= #

df = pd.read_excel(arquivo_excel, sheet_name='Desligados')

# ========= LOOP PRINCIPAL ========= #

for index, row in df.iterrows():

    nome = row['Nome']
    valor = row['Saldo (Atualizado)']
    referencia = row['Mês/Ano']

    # dados adicionais (se existirem)
    endereco = row.get('Endereço', '')
    cep = row.get('CEP', '')
    cidade = row.get('Cidade', '')

    # formatações
    nome_cap = str(nome).title()
    nome_upper = str(nome).upper()

    valor_formatado = formatar_moeda(valor)

    # valor por extenso
    try:
        valor_extenso = num2words(float(valor), lang='pt_BR').upper()
    except:
        valor_extenso = ""

    # data limite (exemplo: 10 dias a partir de hoje)
    data_limite = (hoje.replace(day=hoje.day) + pd.Timedelta(days=10)).strftime('%d/%m/%Y')

    # ========= ABRIR TEMPLATE ========= #

    doc = Document(template_path)

    substituicoes = {
        '{{nome}}': nome_cap,
        '{{nome_upper}}': nome_upper,
        '{{endereco}}': endereco,
        '{{cep}}': cep,
        '{{cidade}}': cidade,
        '{{valor}}': valor_formatado,
        '{{valor_extenso}}': valor_extenso,
        '{{referencia}}': referencia,
        '{{data_limite}}': data_limite,
        '{{dia}}': dia,
        '{{mes}}': mes,
        '{{ano}}': ano
    }

    substituir_texto(doc, substituicoes)

    # ========= SALVAR ========= #

    nome_arquivo = f"{normalize_name(nome)}.docx"
    caminho_saida = os.path.join(pasta_saida, nome_arquivo)

    doc.save(caminho_saida)

print("Cartas geradas com sucesso!")